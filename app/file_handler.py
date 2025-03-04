import time
from tkinter import filedialog, messagebox
import os
import zipfile
import re


class FileHandler:
    @staticmethod
    def is_valid_file_type(file_path):
        valid_extensions = ".bim"
        return file_path.lower().endswith(valid_extensions)

    @staticmethod
    def select_file():
        file_path = filedialog.askopenfilename(
            title="Select a Power BI Model file",
            initialdir=os.path.expanduser("~"),
            filetypes=[("Power BI Model files", "*.bim")],
        )

        return file_path if file_path else ""

    @staticmethod
    def read_file(file_path):
        try:
            with open(file_path, "r") as file:
                return file.read()
        except Exception as e:
            messagebox.showerror("Error", f"Failed to read file: {str(e)}")
            return None

    @staticmethod
    def read_pbix_report(pbix_path):
        MAX_CHARS = 1048546
        MAX_BYTES = 502857

        try:
            with zipfile.ZipFile(pbix_path, "r") as zip_ref:
                contents = {}

                if "Report/Layout" in zip_ref.namelist():
                    with zip_ref.open("Report/Layout", "r") as report_file:
                        content = report_file.read().decode("utf-8", errors="replace")

                        if len(content) > MAX_CHARS:
                            contents["content"] = (
                                content[:MAX_CHARS] + "\n... (content truncated)"
                            )
                        else:
                            contents["content"] = content

                if "DataModel" in zip_ref.namelist():
                    with zip_ref.open("DataModel", "r") as model_file:
                        model_content = model_file.read().decode(
                            "utf-8", errors="replace"
                        )

                        if len(model_content) > MAX_BYTES:
                            contents["modelContent"] = model_content[:MAX_BYTES]
                        else:
                            contents["modelContent"] = model_content

                if not contents:
                    messagebox.showerror(
                        "Error", "No report file found in the PBIX file."
                    )
                    return None

                return contents
        except Exception as e:
            messagebox.showerror("Error", f"Failed to read PBIX file: {str(e)}")
            return None

    @staticmethod
    def save_as_txt(content, filename=None, file_prefix=None):
        """Save content as a text file."""
        if not filename:
            prefix = file_prefix if file_prefix else "analysis"
            filename = os.path.join(
                os.path.expanduser("~"), f"{prefix}_{int(time.time())}.txt"
            )

        try:
            with open(filename, "w", encoding="utf-8") as file:
                file.write(content)
            messagebox.showinfo("Success", f"File saved to {filename}")
            return True
        except Exception as e:
            messagebox.showerror("Error", f"Failed to save text file: {str(e)}")
            return False

    @staticmethod
    def save_as_doc(content, filename=None, file_prefix=None):
        """Save content as a Word document with improved code block formatting."""
        try:
            from docx import Document
            from docx.shared import Pt, RGBColor, Inches
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement
        except ImportError:
            messagebox.showerror(
                "Error",
                "Python-docx package is required to save as DOC. Please install it with 'pip install python-docx'",
            )
            return False

        if not filename:
            prefix = file_prefix if file_prefix else "analysis"
            filename = os.path.join(
                os.path.expanduser("~"), f"{prefix}_{int(time.time())}.docx"
            )

        try:
            doc = Document()
            style = doc.styles["Normal"]
            style.font.name = "Calibri"
            style.font.size = Pt(12)

            # Add logo at the top
            logo_path = os.path.join(
                os.path.dirname(os.path.dirname(__file__)),
                "assets",
                "images",
                "ah_logo.png",
            )

            # Add logo to the document
            logo_paragraph = doc.add_paragraph()
            logo_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            logo_run = logo_paragraph.add_run()
            logo_run.add_picture(logo_path, width=Inches(2.0))

            # Add a space after the logo
            doc.add_paragraph()

            # Define custom heading styles to match our CSS
            h1_style = doc.styles["Heading 1"]
            h1_style.font.name = "Calibri"
            h1_style.font.size = Pt(28)
            h1_style.font.bold = True
            h1_style.font.color.rgb = RGBColor(202, 90, 139)  # #ca5a8b

            h2_style = doc.styles["Heading 2"]
            h2_style.font.name = "Calibri"
            h2_style.font.size = Pt(22)
            h2_style.font.bold = True
            h2_style.font.color.rgb = RGBColor(102, 84, 245)  # #6654f5

            h3_style = doc.styles["Heading 3"]
            h3_style.font.name = "Calibri"
            h3_style.font.size = Pt(18)
            h3_style.font.bold = True
            h3_style.font.color.rgb = RGBColor(102, 84, 245)  # #6654f5

            # Create code block style
            try:
                code_style = doc.styles.add_style("CodeBlock", 1)
                code_style.font.name = "Consolas"
                code_style.font.size = Pt(10)
                code_style.paragraph_format.space_before = Pt(6)
                code_style.paragraph_format.space_after = Pt(6)
                code_style.paragraph_format.keep_together = True
            except:
                # Style might already exist
                code_style = doc.styles["Normal"]

            def process_formatted_text(text, paragraph=None):
                if paragraph is None:
                    paragraph = doc.add_paragraph()

                text = text.replace(r"\*", "____ESCAPED_ASTERISK____")

                # First pass: process text for formatting
                current_position = 0
                remaining_text = text
                found_formatting = False

                # Bold pattern: Match anything between two asterisks
                bold_pattern = re.compile(r"\*\*(.*?)\*\*")

                # Process text for bold formatting
                while True:
                    bold_match = bold_pattern.search(remaining_text)
                    if not bold_match:
                        break

                    # Add text before formatting mark
                    if bold_match.start() > 0:
                        paragraph.add_run(remaining_text[: bold_match.start()])

                    # Add formatted text
                    bold_text = bold_match.group(1)
                    bold_run = paragraph.add_run(bold_text)
                    bold_run.bold = True

                    # Update remaining text
                    remaining_text = remaining_text[bold_match.end() :]
                    found_formatting = True

                # Add any remaining text
                if remaining_text:
                    paragraph.add_run(remaining_text)

                # If no formatting found, just set the text directly
                if not found_formatting and not paragraph.runs:
                    # Replace any temporary placeholders
                    text = text.replace("____ESCAPED_ASTERISK____", "*")
                    paragraph.text = text

                return paragraph

            # Helper function to create and format table
            def create_table_from_markdown(table_lines):
                # Parse the markdown table
                rows = []
                for line in table_lines:
                    # Skip separator lines (---|---|---)
                    if re.match(r"^[\s|]*[-:]+[\s|]*$", line.strip()):
                        continue
                    # Process cells in the row
                    cells = re.findall(r"\|(.*?)(?=\||$)", line + "|")
                    # Clean up cells and remove empty trailing cell if present
                    cells = [cell.strip() for cell in cells]
                    if cells and cells[-1] == "":
                        cells.pop()
                    if cells:  # Only add non-empty rows
                        rows.append(cells)

                if not rows:
                    return None

                # Create the Word table
                table = doc.add_table(rows=len(rows), cols=len(rows[0]))
                table.style = "Table Grid"

                # Format the table
                for i, row in enumerate(rows):
                    for j, cell in enumerate(row):
                        if j < len(table.rows[i].cells):
                            cell_text = cell.strip()
                            cell_paragraph = table.rows[i].cells[j].paragraphs[0]

                            # Use our improved text processing function for cell content
                            process_formatted_text(cell_text, cell_paragraph)

                            # Make header row bold
                            if i == 0:
                                for run in cell_paragraph.runs:
                                    run.bold = True

                return table

            # Helper function to create a code block with syntax highlighting
            def create_code_block(code_lines, language):
                # Create a bordered container for the code block
                code_container = doc.add_paragraph()

                # Add a light gray shaded background for the code block
                shading_element = OxmlElement("w:shd")
                shading_element.set(qn("w:fill"), "F5F5F5")  # Light gray background

                # Add a border to the paragraph
                def set_border(paragraph):
                    p = paragraph._p
                    pPr = p.get_or_add_pPr()
                    pBdr = OxmlElement("w:pBdr")

                    # Add border on all sides
                    for side in ["top", "left", "bottom", "right"]:
                        border = OxmlElement(f"w:{side}")
                        border.set(qn("w:val"), "single")
                        border.set(qn("w:sz"), "4")  # Border width in 1/8 points
                        border.set(qn("w:space"), "0")
                        border.set(qn("w:color"), "CCCCCC")  # Light gray border
                        pBdr.append(border)

                    pPr.append(pBdr)

                    # Add shading
                    pPr.append(shading_element)

                # Create a language label if specified
                if language and language.strip() != "":
                    lang_para = doc.add_paragraph()
                    lang_run = lang_para.add_run(f"{language.strip()}")
                    lang_run.bold = True
                    lang_run.font.size = Pt(9)
                    lang_run.font.color.rgb = RGBColor(70, 70, 70)  # Dark gray

                set_border(code_container)

                # Add the code content with monospace font
                for line in code_lines:
                    code_para = doc.add_paragraph(style="CodeBlock")
                    code_run = code_para.add_run(line)
                    code_run.font.name = "Consolas"  # Monospace font
                    code_run.font.size = Pt(10)

                    # Add paragraph shading
                    set_border(code_para)

                # Add a space after the code block
                doc.add_paragraph()

            # Pre-process content to handle specific patterns
            processed_content = content
            # Replace any escaped asterisks or problematic patterns if needed

            # Split into lines and process
            lines = processed_content.split("\n")
            line_index = 0
            table_lines = []
            in_table = False
            code_block = False
            code_language = ""
            code_lines = []

            while line_index < len(lines):
                line = lines[line_index]

                # Check for code blocks
                if line.startswith("```"):
                    if not code_block:
                        # Start of code block
                        code_block = True
                        code_language = line[3:].strip()  # Extract language identifier
                        code_lines = []
                        line_index += 1
                        continue
                    else:
                        # End of code block
                        code_block = False
                        create_code_block(code_lines, code_language)
                        line_index += 1
                        continue

                # If in code block, collect lines
                if code_block:
                    code_lines.append(line)
                    line_index += 1
                    continue

                # Check if line might be part of a table
                if "|" in line and not line.startswith("```"):
                    # Start collecting table lines
                    if not in_table:
                        in_table = True
                        table_lines = [line]
                    else:
                        table_lines.append(line)
                    line_index += 1
                    continue
                elif in_table:
                    # End of table reached, process it
                    in_table = False
                    create_table_from_markdown(table_lines)
                    doc.add_paragraph()  # Add space after table
                    table_lines = []

                # Process normal line
                if line.startswith("# "):
                    doc.add_heading(line[2:], level=1)
                elif line.startswith("## "):
                    doc.add_heading(line[3:], level=2)
                elif line.startswith("### "):
                    doc.add_heading(line[4:], level=3)
                elif line.startswith("- "):
                    # Process formatted text in list items too
                    paragraph = doc.add_paragraph(style="List Bullet")
                    process_formatted_text(line[2:], paragraph)
                elif line.startswith("* "):
                    paragraph = doc.add_paragraph(style="List Bullet")
                    process_formatted_text(line[2:], paragraph)
                elif line.startswith("1. "):
                    paragraph = doc.add_paragraph(style="List Number")
                    process_formatted_text(line[3:], paragraph)
                elif line.strip() == "":
                    doc.add_paragraph()
                else:
                    # Process the line with our improved formatter
                    process_formatted_text(line)

                line_index += 1

            # Process any remaining table lines
            if table_lines:
                create_table_from_markdown(table_lines)

            # Handle any remaining code block
            if code_block and code_lines:
                create_code_block(code_lines, code_language)

            doc.save(filename)
            messagebox.showinfo("Success", f"File saved to {filename}")
            return True
        except Exception as e:
            messagebox.showerror("Error", f"Failed to save Word document: {str(e)}")
            return False

    @staticmethod
    def save_as_pdf(content, filename=None, file_prefix=None):
        try:
            import markdown2
            import base64
            import xhtml2pdf.pisa as pisa
            from io import BytesIO
            import tempfile
            import shutil
            from xhtml2pdf.default import DEFAULT_FONT
        except ImportError:
            messagebox.showerror(
                "Error",
                "xhtml2pdf and markdown2 packages are required. Install them with 'pip install xhtml2pdf markdown2'",
            )
            return False

        if not filename:
            prefix = file_prefix if file_prefix else "analysis"
            filename = os.path.join(
                os.path.expanduser("~"), f"{prefix}_{int(time.time())}.pdf"
            )

        with tempfile.TemporaryDirectory() as temp_dir:
            try:
                html_body = markdown2.markdown(
                    content,
                    extras=[
                        "fenced-code-blocks",
                        "tables",
                        "code-friendly",
                        "numbering",
                        "cuddled-lists",
                    ],
                )

                logo_path = os.path.join(
                    os.path.dirname(os.path.dirname(__file__)),
                    "assets",
                    "images",
                    "ah_logo.png",
                )

                css_path = os.path.join(
                    os.path.dirname(os.path.dirname(__file__)),
                    "assets",
                    "css",
                    "pdf.css",
                )

                with open(logo_path, "rb") as img_file:
                    logo_data = base64.b64encode(img_file.read()).decode("utf-8")

                fonts_dir = os.path.join(
                    os.path.dirname(os.path.dirname(__file__)), "assets", "fonts"
                )

                temp_fonts_dir = os.path.join(temp_dir, "fonts")
                os.makedirs(temp_fonts_dir, exist_ok=True)

                for font_file in os.listdir(fonts_dir):
                    if font_file.endswith(".ttf"):
                        shutil.copy2(
                            os.path.join(fonts_dir, font_file),
                            os.path.join(temp_fonts_dir, font_file),
                        )

                with open(css_path, "r", encoding="utf-8") as css_file:
                    css_content = css_file.read()

                css_content = css_content.replace("../fonts/", f"{fonts_dir}/")

                styled_html = f"""
                <!DOCTYPE html>
                <html>
                <head>
                    <meta charset="UTF-8">
                    <style>
                        {css_content}
                    </style>
                </head>
                <body>
                    <div id="header_content" class="logo-container">
                        <img src="data:image/png;base64,{logo_data}" class="logo" alt="AnalystHub Logo">
                    </div>
                    {html_body}
                </body>
                </html>
                """

                result_file = open(filename, "w+b")
                pdf_status = pisa.CreatePDF(
                    BytesIO(styled_html.encode("utf-8")), dest=result_file
                )
                result_file.close()

                if not pdf_status.err:
                    messagebox.showinfo("Success", f"File saved to {filename}")
                    return True
                else:
                    messagebox.showerror("Error", "Failed to generate PDF")
                    return False

            except Exception as e:
                import traceback

                error_details = traceback.format_exc()
                messagebox.showerror(
                    "Error", f"Failed to save PDF: {str(e)}\n\nDetails: {error_details}"
                )
                return False

    @staticmethod
    def save_analysis(content, file_format, file_prefix=None):
        """Save analysis content in the specified format."""
        if file_format == "txt":
            return FileHandler.save_as_txt(content, file_prefix=file_prefix)
        elif file_format == "doc":
            return FileHandler.save_as_doc(content, file_prefix=file_prefix)
        elif file_format == "pdf":
            return FileHandler.save_as_pdf(content, file_prefix=file_prefix)
        else:
            messagebox.showerror("Error", f"Unsupported file format: {file_format}")
            return False
